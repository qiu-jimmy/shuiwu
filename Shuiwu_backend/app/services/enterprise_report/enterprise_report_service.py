"""
企业体检报告生成服务
"""
import asyncio
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

import alibabacloud_oss_v2 as oss

from app.agno.teams.enterprise_report_team import EnterpriseReportTeam
from app.infra.logging_config import get_logger

logger = get_logger("app.services.enterprise_report")

# 封面模板路径（相对于项目根目录）
COVER_TEMPLATE_PATH = Path(__file__).parent.parent.parent / "templates" / "cover_template.docx"


def generate_report_number(amount: float) -> str:
    """生成报告编号（使用 UUID，支持并发）

    Args:
        amount: 报价金额（元）

    Returns:
        报告编号，格式：DragonAI-YYYYMM-{S/M/L}-{UUID前8位}
    """
    now = datetime.now()
    year = now.year
    month = now.month

    # 确定单子大小
    if amount < 100000:
        size = "S"
    elif amount < 300000:
        size = "M"
    else:
        size = "L"

    # 生成 UUID 并取前8位作为唯一标识
    unique_id = uuid.uuid4().hex[:8]

    # 生成报告编号
    report_number = f"DragonAI-{year}{month:02d}-{size}-{unique_id}"
    return report_number




class EnterpriseReportService:
    """企业体检报告服务"""

    def __init__(self):
        """初始化服务"""
        # 不在初始化时获取 API Key，而是在使用时获取
        self.api_key = None

    def _get_model_config(self) -> tuple:
        """从模型缓存获取 Qwen 模型配置

        Returns:
            (model_id, api_key, base_url)
        """
        from app.services.models.model_cache import model_cache

        model_id = "qwen-plus"
        model_config = model_cache.get_model_config(model_id)

        if model_config and model_config.get("model_api_key"):
            logger.info(f"从模型缓存获取 {model_id} 配置")
            return (
                model_id,
                model_config.get("model_api_key"),
                model_config.get("model_url")
            )
        else:
            logger.error(f"未在模型缓存中找到 {model_id} 配置")
            raise ValueError(
                f"未配置 {model_id} API Key。"
                f"请在数据库 models 表中添加 {model_id} 模型配置，"
                "或通过 POST /api/models 接口添加。"
            )

    async def generate_report(
            self,
            company_name: str,
            credit_code: str = None,
            user_id: str = "system"
    ) -> dict:
        """生成企业体检报告

        Args:
            company_name: 企业名称
            credit_code: 统一社会信用代码(可选)
            user_id: 用户ID（用于OSS存储路径）

        Returns:
            包含报告信息的字典
        """
        import tempfile
        import os

        # 验证输入
        if not company_name or not company_name.strip():
            raise ValueError("企业名称不能为空")

        company_name = company_name.strip()
        logger.info(f"开始生成企业体检报告: {company_name}")

        # 动态获取模型配置（此时模型缓存应该已经加载）
        model_id, api_key, base_url = self._get_model_config()

        # 生成报告编号（使用 UUID，支持并发）
        report_number = generate_report_number(0.0)  # 不再使用 quote_amount
        logger.info(f"报告编号: {report_number}")

        # 初始化 Team（传递模型配置）
        report_team = EnterpriseReportTeam(model_id=model_id, api_key=api_key, base_url=base_url)

        # 调用企查查 API 获取企业信息
        logger.info("正在查询企业基本信息...")
        company_info_text = await report_team.fetch_company_info(
            company_name=company_name,
            credit_code=credit_code
        )

        # 准备数据
        data = {
            'report_number': report_number,
            'date': datetime.now().strftime("%Y年%m月%d日"),
            'project_name': company_name,
        }

        # 生成实施周期
        logger.info("正在生成改进计划与持续管理...")
        timeline = await report_team.generate_implementation_timeline(company_info_text)

        # 生成完整Markdown文档
        logger.info("正在并行生成所有内容...")
        markdown_content = await report_team.generate_all_content_with_titles(
            company_info_text, company_name, timeline, 0.0
        )
        logger.info("完整 Markdown 文档生成完成")

        # 生成文件名
        safe_project_name = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_project_name = safe_project_name.replace(' ', '_')
        filename = f"{safe_project_name}_{report_number}.docx"

        # 创建临时文件路径（使用系统临时目录）
        temp_file_path = os.path.join(tempfile.gettempdir(), f"temp_report_{report_number}.docx")

        # 转换为Word文档（保存到临时文件）
        logger.info("正在生成Word文档...")
        try:
            await self._generate_word_document(
                markdown_content,
                temp_file_path,
                company_name,
                report_number,
                data.get('date', '')
            )
            logger.info("Word文档生成完成")
        except Exception as e:
            logger.error(f"Word文档生成失败: {e}")
            raise

        # 上传到OSS
        logger.info("正在上传文档到OSS...")
        try:
            file_url = await self._upload_to_oss(
                user_id=user_id,
                file_path=temp_file_path,
                filename=filename
            )
            logger.info(f"文档上传成功: {file_url}")
        except Exception as e:
            logger.error(f"上传文档到OSS失败: {e}")
            raise
        finally:
            # 删除临时文件
            try:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    logger.info(f"临时文件已删除: {temp_file_path}")
            except Exception as e:
                logger.warning(f"删除临时文件失败: {e}")

        return {
            'report_number': report_number,
            'file_url': file_url,
            'file_name': filename,
            'project_name': company_name,
            'generated_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }

    async def _generate_word_document(
            self,
            markdown_content: str,
            output_path: str,
            project_name: str,
            report_number: str,
            date: str
    ):
        """生成Word文档

        Args:
            markdown_content: Markdown内容
            output_path: 输出路径
            project_name: 项目名称
            report_number: 报告编号
            date: 日期
        """
        # 检查 Pandoc 是否可用
        pandoc_available = await self._check_pandoc()

        if pandoc_available:
            # 使用 Pandoc 转换（推荐方式）
            logger.info("使用 Pandoc 生成 Word 文档")
            await self._convert_with_pandoc(markdown_content, output_path, project_name, report_number, date)
        else:
            # 使用 python-docx 直接生成（备用方式）
            logger.warning("Pandoc 不可用，使用 python-docx 生成文档（格式可能不够美观）")
            await self._convert_with_python_docx(markdown_content, output_path, project_name, report_number, date)

    async def _upload_to_oss(self, user_id: str, file_path: str, filename: str) -> str:
        """上传文件到OSS

        Args:
            user_id: 用户ID
            file_path: 本地文件路径
            filename: 文件名

        Returns:
            OSS文件URL
        """
        from app.infra.oss_client import oss_client_manager

        # 确保OSS客户端已初始化
        if not oss_client_manager.is_initialized():
            if not oss_client_manager.initialize_from_db():
                raise Exception("OSS客户端未初始化，请先配置OSS环境变量或通过API配置")

        # 生成OSS存储路径：enterprise_reports/{user_id}/{year}/{month}/{filename}
        now = datetime.now()
        oss_key = f"enterprise_reports/{user_id}/{now.year}/{now.month:02d}/{filename}"

        # 读取文件内容
        with open(file_path, 'rb') as f:
            file_content = f.read()

        # 上传到OSS
        client = oss_client_manager.client
        bucket = oss_client_manager.bucket

        result = client.put_object(oss.PutObjectRequest(
            bucket=bucket,
            key=oss_key,
            body=file_content,
        ))

        if result.status_code != 200:
            raise Exception(f"OSS上传失败，状态码: {result.status_code}")

        # 构建文件URL
        region = oss_client_manager._config.get('region', 'cn-hangzhou')
        endpoint = oss_client_manager._config.get('endpoint')

        if endpoint and endpoint.strip():
            # 使用自定义endpoint（去除协议前缀）
            endpoint_clean = endpoint.replace('https://', '').replace('http://', '').strip()
            file_url = f"https://{bucket}.{endpoint_clean}/{oss_key}"
        else:
            # 使用标准endpoint格式：bucket.oss-region.aliyuncs.com
            file_url = f"https://{bucket}.oss-{region}.aliyuncs.com/{oss_key}"

        logger.info(f"文件已上传到OSS: {oss_key}")
        return file_url

    async def _check_pandoc(self) -> bool:
        """检查 Pandoc 是否可用"""
        import subprocess
        try:
            result = await asyncio.to_thread(
                subprocess.run,
                ['pandoc', '--version'],
                capture_output=True,
                text=True,
                timeout=5
            )
            is_available = result.returncode == 0
            if is_available:
                logger.info(f"Pandoc 可用: {result.stdout.split()[1]}")
            return is_available
        except FileNotFoundError:
            logger.warning("未找到 Pandoc，请安装 Pandoc 以获得更好的文档格式")
            return False
        except Exception as e:
            logger.warning(f"检查 Pandoc 时出错: {e}")
            return False

    async def _convert_with_pandoc(self, markdown_content: str, output_path: str, company_name: str, report_number: str, report_date: str):
        """使用 Pandoc 转换

        Args:
            markdown_content: Markdown内容
            output_path: 输出文件路径
            company_name: 公司名称，用于水印和页眉
            report_number: 报告编号
            report_date: 报告日期
        """
        import subprocess
        import tempfile
        import os
        import uuid

        # 删除已存在的文件
        if Path(output_path).exists():
            Path(output_path).unlink()

        # 创建唯一的临时 Markdown 文件名（使用 UUID 避免并发冲突）
        temp_filename = f"temp_markdown_{uuid.uuid4().hex}.md"
        md_path = str(Path(tempfile.gettempdir()) / temp_filename)

        try:
            # 写入临时文件
            with open(md_path, 'w', encoding='utf-8') as md_file:
                md_file.write(markdown_content)

            # 使用 Pandoc 转换
            logger.info(f"使用 Pandoc 转换 Markdown 到 Word: {md_path} -> {output_path}")
            result = await asyncio.to_thread(
                subprocess.run,
                ['pandoc', '-f', 'markdown', '-t', 'docx', '-o', output_path, md_path],
                capture_output=True,
                text=True,
                check=True,
                timeout=60
            )
            logger.info(f"Pandoc 转换完成: {output_path}")

            # 后处理：设置文档格式（传递公司名称、报告编号和日期）
            await self._postprocess_document(output_path, company_name, report_number, report_date)

        except subprocess.CalledProcessError as e:
            logger.error(f"Pandoc 转换失败: {e.stderr}")
            raise RuntimeError(f"Pandoc 转换失败: {e.stderr}")
        except Exception as e:
            logger.error(f"Pandoc 转换异常: {e}")
            raise
        finally:
            # 删除临时文件
            try:
                if Path(md_path).exists():
                    os.unlink(md_path)
            except Exception as e:
                logger.warning(f"删除临时文件失败: {e}")

    async def _postprocess_document(self, output_path: str, company_name: str = "税小通", report_number: str = "", report_date: str = ""):
        """后处理生成的 Word 文档，设置格式（与原始脚本保持一致）

        Args:
            output_path: 文档路径
            company_name: 公司名称，用于水印和页眉
            report_number: 报告编号
            report_date: 报告日期
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            logger.info(f"后处理 Word 文档: {output_path}")
            doc = Document(output_path)

            # 设置正文字体：宋体小四（12pt）
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            doc.styles['Normal'].font.size = Pt(12)

            # 设置标题样式字体（标题使用微软雅黑，正文使用宋体）
            # 设置所有标题级别为黑色，微软雅黑字体
            for i in range(1, 10):
                heading_style = f'Heading {i}'
                if heading_style in doc.styles:
                    doc.styles[heading_style].font.name = '微软雅黑'
                    doc.styles[heading_style]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    doc.styles[heading_style].font.color.rgb = RGBColor(0, 0, 0)  # 黑色

            # 一级标题：微软雅黑17.5pt，黑色
            if 'Heading 1' in doc.styles:
                doc.styles['Heading 1'].font.size = Pt(17.5)

            # 二级标题：微软雅黑三号（16pt），黑色
            if 'Heading 2' in doc.styles:
                doc.styles['Heading 2'].font.size = Pt(16)

            # 三级标题：微软雅黑四号（14pt），黑色
            if 'Heading 3' in doc.styles:
                doc.styles['Heading 3'].font.size = Pt(14)

            # 遍历所有段落，设置格式和字体
            for para in doc.paragraphs:
                # 设置段落格式：1.5倍行距，段后间距6pt
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(6)

                # 确保所有run都使用正确的字体和颜色
                for run in para.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    if run.font.size is None:
                        run.font.size = Pt(12)
                    # 确保所有文字都是黑色
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    # 强制设置XML颜色属性（覆盖任何现有颜色）
                    rpr = run._element.get_or_add_rPr()
                    # 移除现有的颜色元素
                    for color_elem in rpr.xpath('.//w:color'):
                        rpr.remove(color_elem)
                    # 添加黑色颜色
                    color_elem = OxmlElement('w:color')
                    color_elem.set(qn('w:val'), '000000')
                    rpr.append(color_elem)

            # 添加封面页、水印、页眉页脚和页码
            if report_number and report_date:
                self._add_cover_page(doc, company_name, report_number, report_date)
            self._add_watermark(doc, company_name)
            self._add_header_footer(doc, company_name)
            self._add_page_numbers(doc)

            # 保存修改后的文档
            doc.save(output_path)
            logger.info("Word 文档后处理完成")

        except ImportError:
            logger.warning("python-docx 未安装，跳过后处理步骤")
        except Exception as e:
            logger.warning(f"后处理文档失败: {e}，但文档已生成")

    def _add_cover_page(self, doc, company_name: str, report_number: str, report_date: str):
        """在文档开头添加封面页

        Args:
            doc: Document对象
            company_name: 公司名称
            report_number: 报告编号
            report_date: 报告日期
        """
        try:
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            logger.info("开始添加封面页")

            # 获取文档的第一个段落（用于在其前插入内容）
            body_element = doc._element.body

            # 创建段落元素
            # 标题段落
            title = OxmlElement('w:p')
            title_pr = OxmlElement('w:pPr')
            titlejc = OxmlElement('w:jc')
            titlejc.set(qn('w:val'), 'center')
            title_pr.append(titlejc)
            title.append(title_pr)

            title_r = OxmlElement('w:r')
            title_pr_r = OxmlElement('w:rPr')
            title_sz = OxmlElement('w:sz')
            title_sz.set(qn('w:val'), '56')  # 28pt
            title_b = OxmlElement('w:b')
            title_b.set(qn('w:val'), '1')
            title_rfonts = OxmlElement('w:rFonts')
            title_rfonts.set(qn('w:eastAsia'), '微软雅黑')
            title_pr_r.append(title_sz)
            title_pr_r.append(title_b)
            title_pr_r.append(title_rfonts)
            title_r.append(title_pr_r)
            title_t = OxmlElement('w:t')
            title_t.text = company_name
            title_r.append(title_t)
            title.append(title_r)

            # 段后间距
            titlespacing = OxmlElement('w:spacing')
            titlespacing.set(qn('w:after'), '240')
            title_pr.append(titlespacing)

            # 副标题
            subtitle = OxmlElement('w:p')
            sub_pr = OxmlElement('w:pPr')
            subjc = OxmlElement('w:jc')
            subjc.set(qn('w:val'), 'center')
            sub_pr.append(subjc)
            subtitle.append(sub_pr)

            sub_r = OxmlElement('w:r')
            sub_pr_r = OxmlElement('w:rPr')
            sub_sz = OxmlElement('w:sz')
            sub_sz.set(qn('w:val'), '44')  # 22pt
            sub_b = OxmlElement('w:b')
            sub_b.set(qn('w:val'), '1')
            sub_rfonts = OxmlElement('w:rFonts')
            sub_rfonts.set(qn('w:eastAsia'), '微软雅黑')
            sub_pr_r.append(sub_sz)
            sub_pr_r.append(sub_b)
            sub_pr_r.append(sub_rfonts)
            sub_r.append(sub_pr_r)
            sub_t = OxmlElement('w:t')
            sub_t.text = "企业体检报告"
            sub_r.append(sub_t)
            subtitle.append(sub_r)

            # 段后间距
            subspacing = OxmlElement('w:spacing')
            subspacing.set(qn('w:after'), '1600')
            sub_pr.append(subspacing)

            # 报告编号
            rn_para = OxmlElement('w:p')
            rn_pr = OxmlElement('w:pPr')
            rn_jc = OxmlElement('w:jc')
            rn_jc.set(qn('w:val'), 'center')
            rn_pr.append(rn_jc)
            rn_para.append(rn_pr)

            rn_r = OxmlElement('w:r')
            rn_pr_r = OxmlElement('w:rPr')
            rn_sz = OxmlElement('w:sz')
            rn_sz.set(qn('w:val'), '28')  # 14pt
            rn_rfonts = OxmlElement('w:rFonts')
            rn_rfonts.set(qn('w:eastAsia'), '微软雅黑')
            rn_color = OxmlElement('w:color')
            rn_color.set(qn('w:val'), '505050')
            rn_pr_r.append(rn_sz)
            rn_pr_r.append(rn_rfonts)
            rn_pr_r.append(rn_color)
            rn_r.append(rn_pr_r)
            rn_t = OxmlElement('w:t')
            rn_t.text = f"报告编号：{report_number}"
            rn_r.append(rn_t)
            rn_para.append(rn_r)

            # 段后间距
            rnspacing = OxmlElement('w:spacing')
            rnspacing.set(qn('w:after'), '240')
            rn_pr.append(rnspacing)

            # 日期
            date_para = OxmlElement('w:p')
            date_pr = OxmlElement('w:pPr')
            date_jc = OxmlElement('w:jc')
            date_jc.set(qn('w:val'), 'center')
            date_pr.append(date_jc)
            date_para.append(date_pr)

            date_r = OxmlElement('w:r')
            date_pr_r = OxmlElement('w:rPr')
            date_sz = OxmlElement('w:sz')
            date_sz.set(qn('w:val'), '28')  # 14pt
            date_rfonts = OxmlElement('w:rFonts')
            date_rfonts.set(qn('w:eastAsia'), '微软雅黑')
            date_color = OxmlElement('w:color')
            date_color.set(qn('w:val'), '505050')
            date_pr_r.append(date_sz)
            date_pr_r.append(date_rfonts)
            date_pr_r.append(date_color)
            date_r.append(date_pr_r)
            date_t = OxmlElement('w:t')
            date_t.text = f"生成日期：{report_date}"
            date_r.append(date_t)
            date_para.append(date_r)

            # 段后间距
            datespacing = OxmlElement('w:spacing')
            datespacing.set(qn('w:after'), '2400')
            date_pr.append(datespacing)

            # 底部信息
            footer_para = OxmlElement('w:p')
            footer_pr = OxmlElement('w:pPr')
            footer_jc = OxmlElement('w:jc')
            footer_jc.set(qn('w:val'), 'center')
            footer_pr.append(footer_jc)
            footer_para.append(footer_pr)

            footer_r = OxmlElement('w:r')
            footer_pr_r = OxmlElement('w:rPr')
            footer_sz = OxmlElement('w:sz')
            footer_sz.set(qn('w:val'), '24')  # 12pt
            footer_rfonts = OxmlElement('w:rFonts')
            footer_rfonts.set(qn('w:eastAsia'), '微软雅黑')
            footer_color = OxmlElement('w:color')
            footer_color.set(qn('w:val'), '808080')
            footer_pr_r.append(footer_sz)
            footer_pr_r.append(footer_rfonts)
            footer_pr_r.append(footer_color)
            footer_r.append(footer_pr_r)
            footer_t = OxmlElement('w:t')
            footer_t.text = "税小通 · 企业体检报告"
            footer_r.append(footer_t)
            footer_para.append(footer_r)

            # 分页符
            page_break = OxmlElement('w:p')
            pb_pr = OxmlElement('w:pPr')
            page_break.append(pb_pr)
            pb_r = OxmlElement('w:r')
            pb_br = OxmlElement('w:br')
            pb_br.set(qn('w:type'), 'page')
            pb_r.append(pb_br)
            page_break.append(pb_r)

            # 在文档开头插入所有元素（按相反顺序）
            body_element.insert(0, page_break)
            body_element.insert(0, footer_para)
            body_element.insert(0, date_para)
            body_element.insert(0, rn_para)
            body_element.insert(0, subtitle)
            body_element.insert(0, title)

            logger.info("封面页添加完成")

        except Exception as e:
            logger.warning(f"添加封面页失败: {e}")

    def _add_watermark(self, doc, company_name: str):
        """添加水印

        Args:
            doc: Document对象
            company_name: 公司名称（水印文字）
        """
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            # 获取文档的section
            for section in doc.sections:
                # 获取header部分，水印通常放在header中
                header = section.header

                # 创建新的段落用于放置水印
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

                # 清空现有内容
                paragraph.clear()

                # 创建运行对象
                run = paragraph.add_run()

                # 设置水印属性
                run.text = company_name

                # 设置字体格式
                run.font.name = '微软雅黑'
                run.font.size = Pt(48)
                run.font.color.rgb = RGBColor(200, 200, 200)  # 浅灰色
                run.font.bold = True

                # 设置段落格式
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            logger.info(f"水印添加完成: {company_name}")
        except Exception as e:
            logger.warning(f"添加水印失败: {e}")

    def _add_header_footer(self, doc, company_name: str):
        """添加页眉页脚

        Args:
            doc: Document对象
            company_name: 公司名称
        """
        try:
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from datetime import datetime

            for section in doc.sections:
                # 添加页眉（左侧显示公司名称，右侧显示报告日期）
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

                # 清空默认内容
                header_para.clear()

                # 添加左侧文本（公司名称）
                run1 = header_para.add_run(f"{company_name}")
                run1.font.name = '微软雅黑'
                run1.font.size = Pt(10)
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

                # 添加制表符
                header_para.add_run("\t")

                # 添加右侧文本（日期）
                from datetime import datetime
                run2 = header_para.add_run(f"{datetime.now().strftime('%Y年%m月%d日')}")
                run2.font.name = '微软雅黑'
                run2.font.size = Pt(10)
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

                # 设置页眉段落格式
                header_para.paragraph_format.line_spacing = 1

                # 添加页脚（居中显示版权信息）
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.clear()

                footer_text = footer_para.add_run("税小通 - 企业体检报告 | 机密文件")
                footer_text.font.name = '微软雅黑'
                footer_text.font.size = Pt(9)
                footer_text._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                footer_text.font.color.rgb = RGBColor(128, 128, 128)  # 灰色

                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                footer_para.paragraph_format.line_spacing = 1

            logger.info("页眉页脚添加完成")
        except Exception as e:
            logger.warning(f"添加页眉页脚失败: {e}")

    def _add_page_numbers(self, doc):
        """添加页码

        Args:
            doc: Document对象
        """
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            for section in doc.sections:
                footer = section.footer

                # 获取或创建页脚段落
                if len(footer.paragraphs) == 0:
                    footer_para = footer.add_paragraph()
                else:
                    footer_para = footer.paragraphs[0]

                # 清空现有内容并添加页码
                footer_para.clear()
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                # 添加页码文本和域代码
                run = footer_para.add_run("第 ")
                run.font.name = '微软雅黑'
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

                # 添加页码域
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')

                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"

                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')

                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)

                # 添加"页"文字
                run2 = footer_para.add_run(" 页")
                run2.font.name = '微软雅黑'
                run2.font.size = Pt(9)
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            logger.info("页码添加完成")
        except Exception as e:
            logger.warning(f"添加页码失败: {e}")

    async def _convert_with_python_docx(
            self,
            markdown_content: str,
            output_path: str,
            project_name: str,
            report_number: str,
            date: str
    ):
        """使用 python-docx 生成（简化版）"""
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # 添加封面页（使用项目名称作为公司名称）
        self._add_cover_page(doc, project_name, report_number, date)

        # 解析 Markdown 并添加内容（简化处理）
        lines = markdown_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('# '):
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        # 添加水印、页眉页脚和页码
        self._add_watermark(doc, project_name)
        self._add_header_footer(doc, project_name)
        self._add_page_numbers(doc)

        doc.save(output_path)


# 创建全局服务实例
enterprise_report_service = EnterpriseReportService()
