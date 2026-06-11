"""
Markdown 处理工具
"""
import re
import tempfile
import os
from docx import Document


class MarkdownHandler:
    """Markdown 处理工具类"""
    
    @staticmethod
    def ensure_pandoc():
        """确保 Pandoc 可用"""
        try:
            import pypandoc
            # 尝试获取版本，如果成功说明 Pandoc 可用
            pypandoc.get_pandoc_version()
            return True
        except OSError:
            # Pandoc 不可用，尝试下载
            try:
                from pypandoc.pandoc_download import download_pandoc
                download_pandoc()
                return True
            except Exception:
                return False
        except ImportError:
            return False
    
    @staticmethod
    def clean_markdown_simple(markdown_text: str) -> str:
        """简单清理 Markdown（移除格式标记）"""
        cleaned_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', markdown_text)
        cleaned_text = re.sub(r'^#{1,6}\s+', '', cleaned_text, flags=re.MULTILINE)
        return cleaned_text
    
    @staticmethod
    def convert_markdown_to_docx_via_pandoc(markdown_text: str, output_path: str) -> Document:
        """使用 Pandoc 将 Markdown 转换为 docx 文档
        
        Args:
            markdown_text: Markdown 格式的文本内容
            output_path: 输出文件路径
            
        Returns:
            转换后的 Document 对象
        """
        if not MarkdownHandler.ensure_pandoc():
            raise RuntimeError("Pandoc 不可用，请安装 Pandoc")
        
        try:
            import pypandoc
            
            # 预处理 Markdown：处理可能被误判为 YAML metadata 的内容
            processed_markdown = markdown_text
            
            # 1. 如果文档以 `---` 开头，在前面添加内容，避免被误判为 YAML metadata
            if processed_markdown.strip().startswith('---'):
                processed_markdown = '# 文档\n\n' + processed_markdown
            
            # 2. 转义文档中所有独立的 `---` 行（可能是分隔线），避免被误判为 YAML metadata
            # 将独立的 `---` 行替换为 `***`（Markdown 分隔线）
            lines = processed_markdown.split('\n')
            processed_lines = []
            for i, line in enumerate(lines):
                stripped = line.strip()
                # 如果这一行只有 `---`（可能前后有空格），替换为 `***`
                if stripped == '---':
                    # 保持原有的缩进，但将内容替换为 `***`
                    indent = len(line) - len(line.lstrip())
                    processed_lines.append(' ' * indent + '***')
                else:
                    processed_lines.append(line)
            processed_markdown = '\n'.join(processed_lines)
            
            # 创建临时 Markdown 文件
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp_md:
                tmp_md.write(processed_markdown)
                tmp_md_path = tmp_md.name
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                tmp_docx_path = tmp_docx.name
            
            try:
                # 使用 Pandoc 将 Markdown 转换为 docx
                # 使用 gfm (GitHub Flavored Markdown) 格式，它不包含 YAML metadata block
                pypandoc.convert_file(
                    tmp_md_path,
                    'docx',
                    outputfile=tmp_docx_path,
                    format='gfm',  # GitHub Flavored Markdown，不包含 YAML metadata
                    extra_args=[
                        '--standalone',
                        '--wrap=none',
                    ]
                )
                
                # 读取转换后的 docx 文件
                doc = Document(tmp_docx_path)
                return doc
                
            finally:
                # 清理临时文件
                try:
                    os.unlink(tmp_md_path)
                except:
                    pass
                try:
                    os.unlink(tmp_docx_path)
                except:
                    pass
                    
        except ImportError:
            raise ImportError("需要安装 pypandoc 库。请运行: pip install pypandoc")
        except Exception as e:
            raise RuntimeError(f"使用 Pandoc 转换失败: {str(e)}")
    
    @staticmethod
    def add_markdown_to_document_via_pandoc(doc: Document, markdown_text: str):
        """使用 Pandoc 将 Markdown 内容插入到文档中
        
        Args:
            doc: 目标文档对象
            markdown_text: Markdown 格式的文本内容
        """
        if not markdown_text or not markdown_text.strip():
            return
        
        # 检查 Pandoc 是否可用
        if not MarkdownHandler.ensure_pandoc():
            # 如果 Pandoc 不可用，回退到清理模式
            cleaned_text = MarkdownHandler.clean_markdown_simple(markdown_text)
            paragraphs = [p.strip() for p in cleaned_text.split('\n') if p.strip()]
            for para_text in paragraphs:
                doc.add_paragraph(f"　　{para_text.lstrip()}")
            return
        
        try:
            import pypandoc
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp_md:
                tmp_md.write(markdown_text)
                tmp_md_path = tmp_md.name
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                tmp_docx_path = tmp_docx.name
            
            try:
                # 使用 Pandoc 将 Markdown 转换为 docx
                # 使用 gfm (GitHub Flavored Markdown) 格式，它不包含 YAML metadata block
                pypandoc.convert_file(
                    tmp_md_path,
                    'docx',
                    outputfile=tmp_docx_path,
                    format='gfm',  # GitHub Flavored Markdown，不包含 YAML metadata
                    extra_args=[
                        '--standalone',
                        '--wrap=none',
                    ]
                )
                
                # 读取转换后的 docx 文件
                temp_doc = Document(tmp_docx_path)
                
                # 将临时文档的内容复制到当前文档
                for para in temp_doc.paragraphs:
                    if para.text.strip():  # 只复制非空段落
                        new_para = doc.add_paragraph()
                        # 复制段落格式
                        new_para.style = para.style
                        # 复制所有 runs（包括格式）
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.font.name = run.font.name
                            new_run.font.size = run.font.size
                            new_run.font.bold = run.font.bold
                            new_run.font.italic = run.font.italic
                            if run._element.rPr.rFonts is not None:
                                new_run._element.rPr.rFonts = run._element.rPr.rFonts
                
                # 复制表格
                for table in temp_doc.tables:
                    # 创建新表格
                    new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
                
            finally:
                # 清理临时文件
                try:
                    os.unlink(tmp_md_path)
                except:
                    pass
                try:
                    os.unlink(tmp_docx_path)
                except:
                    pass
                    
        except ImportError:
            # 如果 pypandoc 未安装，回退到清理模式
            cleaned_text = MarkdownHandler.clean_markdown_simple(markdown_text)
            paragraphs = [p.strip() for p in cleaned_text.split('\n') if p.strip()]
            for para_text in paragraphs:
                doc.add_paragraph(f"　　{para_text.lstrip()}")
        except Exception as e:
            # 如果转换失败，回退到清理模式
            cleaned_text = MarkdownHandler.clean_markdown_simple(markdown_text)
            paragraphs = [p.strip() for p in cleaned_text.split('\n') if p.strip()]
            for para_text in paragraphs:
                doc.add_paragraph(f"　　{para_text.lstrip()}")

