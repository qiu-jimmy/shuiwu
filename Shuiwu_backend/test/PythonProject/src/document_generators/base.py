"""
文档生成器基础类和工具函数
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from pathlib import Path


class BaseDocumentGenerator:
    """文档生成器基类"""
    
    def __init__(self):
        """初始化文档生成器"""
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档基本格式"""
        # 设置正文字体：微软雅黑小四（12pt），英文Times New Roman
        self.doc.styles['Normal'].font.name = 'Times New Roman'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        self.doc.styles['Normal'].font.size = Pt(12)
        
        # 设置标题样式字体
        # 一级标题：微软雅黑17.5pt
        if 'Heading 1' in self.doc.styles:
            self.doc.styles['Heading 1'].font.name = 'Times New Roman'
            self.doc.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            self.doc.styles['Heading 1'].font.size = Pt(17.5)
        
        # 二级标题：微软雅黑三号（16pt）
        if 'Heading 2' in self.doc.styles:
            self.doc.styles['Heading 2'].font.name = 'Times New Roman'
            self.doc.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            self.doc.styles['Heading 2'].font.size = Pt(16)
        
        # 三级标题：微软雅黑四号（14pt）
        if 'Heading 3' in self.doc.styles:
            self.doc.styles['Heading 3'].font.name = 'Times New Roman'
            self.doc.styles['Heading 3']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            self.doc.styles['Heading 3'].font.size = Pt(14)
    
    def _set_paragraph_format(self, para):
        """统一设置段落格式（行间距、段间距等）"""
        para_format = para.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # 多倍行距
        para_format.line_spacing = 1.5  # 1.5倍行距
        para_format.space_before = Pt(0)  # 段前间距
        para_format.space_after = Pt(6)  # 段后间距（6磅）
        para_format.first_line_indent = Pt(0)  # 首行缩进（0，不缩进）
    
    def add_title(self, title: str):
        """添加标题（微软雅黑17.5pt）"""
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.font.size = Pt(17.5)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(12)
    
    def add_heading(self, text: str, level: int = 1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        
        # 设置字体大小
        if level == 1:
            font_size = Pt(17.5)  # 一级标题：微软雅黑17.5pt
        elif level == 2:
            font_size = Pt(16)  # 二级标题：微软雅黑三号（16pt）
        elif level == 3:
            font_size = Pt(14)  # 三级标题：微软雅黑四号（14pt）
        else:
            font_size = Pt(12)  # 默认小四
        
        # 设置样式字体
        heading.style.font.name = 'Times New Roman'
        heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        heading.style.font.size = font_size
        
        # 确保所有run都使用正确的字体
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = font_size
    
    def add_paragraph(self, text: str):
        """添加段落（微软雅黑小四，英文Times New Roman，统一行间距和段间距）"""
        para = self.doc.add_paragraph(text)
        para.style.font.name = 'Times New Roman'
        para.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        para.style.font.size = Pt(12)
        
        # 设置段落格式：行间距和段间距
        self._set_paragraph_format(para)
        
        # 确保所有run都使用正确的字体
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(12)
    
    def add_paragraph_compact(self, text: str):
        """添加紧凑段落：行距统一，但段前/段后不额外加间距，适合连续正文"""
        para = self.doc.add_paragraph(text)
        para.style.font.name = 'Times New Roman'
        para.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        para.style.font.size = Pt(12)

        # 基于通用格式再微调：去掉段前/段后间距
        self._set_paragraph_format(para)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        for run in para.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(12)
    
    def _setup_document_for_existing_doc(self, doc: Document):
        """为已存在的文档设置基本格式"""
        # 设置正文字体：微软雅黑小四（12pt），英文Times New Roman
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.styles['Normal'].font.size = Pt(12)
        
        # 设置标题样式字体
        # 一级标题：微软雅黑17.5pt
        if 'Heading 1' in doc.styles:
            doc.styles['Heading 1'].font.name = 'Times New Roman'
            doc.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            doc.styles['Heading 1'].font.size = Pt(17.5)
        
        # 二级标题：微软雅黑三号（16pt）
        if 'Heading 2' in doc.styles:
            doc.styles['Heading 2'].font.name = 'Times New Roman'
            doc.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            doc.styles['Heading 2'].font.size = Pt(16)
        
        # 三级标题：微软雅黑四号（14pt）
        if 'Heading 3' in doc.styles:
            doc.styles['Heading 3'].font.name = 'Times New Roman'
            doc.styles['Heading 3']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            doc.styles['Heading 3'].font.size = Pt(14)
        
        # 遍历所有段落，确保字体设置正确
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name != 'Times New Roman':
                    run.font.name = 'Times New Roman'
                if run._element.rPr.rFonts is None:
                    run._element.rPr.rFonts = run._element.rPr.rFonts.__class__()
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                if run.font.size is None:
                    run.font.size = Pt(12)

